"""
文档处理模块 - 支持图片 OCR、PDF、Word 文档识别
使用 Kimi (Moonshot) API 进行图片识别
"""

import base64
import io
from pathlib import Path
from typing import Union
from openai import OpenAI
from config import settings

# 文档处理库
try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# 支持的文件类型
SUPPORTED_IMAGE_TYPES = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'}
SUPPORTED_DOCUMENT_TYPES = {'.pdf', '.docx', '.doc', '.txt', '.xlsx', '.xls'}
SUPPORTED_VIDEO_TYPES = {'.mp4', '.mov', '.avi', '.mkv', '.webm'}


class DocumentProcessor:
    """文档处理器"""

    def __init__(self):
        # 使用 OpenAI 兼容格式的 Kimi API
        self.client = OpenAI(
            api_key=settings.kimi_api_key,
            base_url=settings.kimi_base_url
        )

    def process_file(self, file_content: bytes, filename: str) -> dict:
        """
        处理上传的文件

        Args:
            file_content: 文件二进制内容
            filename: 文件名（用于判断类型）

        Returns:
            dict: {
                'type': 'image' | 'document' | 'unsupported',
                'content': 提取的文本内容,
                'images': [图片base64列表, 仅图片类型时有],
                'filename': 文件名
            }
        """
        ext = Path(filename).suffix.lower()

        if ext in SUPPORTED_IMAGE_TYPES:
            return self._process_image(file_content, filename)
        elif ext in SUPPORTED_VIDEO_TYPES:
            return self._process_video(file_content, filename)
        elif ext == '.pdf':
            return self._process_pdf(file_content, filename)
        elif ext in ['.docx', '.doc']:
            return self._process_word(file_content, filename)
        elif ext in ['.xlsx', '.xls']:
            return self._process_excel(file_content, filename)
        elif ext == '.txt':
            return self._process_text(file_content, filename)
        else:
            return {
                'type': 'unsupported',
                'content': f'不支持的文件类型: {ext}',
                'filename': filename
            }

    def _process_image(self, file_content: bytes, filename: str) -> dict:
        """处理图片 - 使用 Kimi 视觉能力识别"""
        try:
            # 压缩图片以节省 token
            image = Image.open(io.BytesIO(file_content))

            # 转换为 RGB (处理RGBA等格式)
            if image.mode in ('RGBA', 'P'):
                image = image.convert('RGB')

            # 压缩大图片
            max_size = (1024, 1024)
            image.thumbnail(max_size, Image.Resampling.LANCZOS)

            # 转换为 base64
            buffer = io.BytesIO()
            image.save(buffer, format='JPEG', quality=85)
            base64_image = base64.b64encode(buffer.getvalue()).decode()

            # 使用 Kimi 识别图片内容
            extracted_text = self._ocr_with_kimi(base64_image)

            return {
                'type': 'image',
                'content': extracted_text,
                'images': [base64_image],
                'filename': filename
            }

        except Exception as e:
            return {
                'type': 'image',
                'content': f'图片处理失败: {str(e)}',
                'filename': filename
            }

    def _ocr_with_kimi(self, base64_image: str) -> str:
        """使用 Kimi 视觉模型识别图片中的文字"""
        try:
            response = self.client.chat.completions.create(
                model=settings.kimi_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                }
                            },
                            {
                                "type": "text",
                                "text": """请识别图片中的文字内容，并严格按照以下JSON格式输出，不要添加任何markdown标记或其他说明文字：

如果是聊天记录截图：
{"type": "chat", "time": "识别到的时间或空字符串", "participants": ["参与者A", "参与者B"], "messages": [{"sender": "A", "content": "消息内容"}, {"sender": "B", "content": "消息内容"}], "summary": "对话要点总结，不超过50字"}

如果是文档/表格：
{"type": "document", "title": "文档标题或空字符串", "sections": [{"heading": "段落标题", "content": "段落内容"}, {"heading": "段落标题", "content": "段落内容"}], "key_points": ["要点1", "要点2", "要点3"]}

如果是名片/联系方式：
{"type": "contact", "name": "姓名或空字符串", "company": "公司或空字符串", "title": "职位或空字符串", "phone": "电话或空字符串", "email": "邮箱或空字符串", "address": "地址或空字符串", "other": "其他信息或空字符串"}

重要提示：
1. 必须返回纯JSON格式，不要添加```json标记
2. 不要添加任何解释说明文字
3. 所有字段必须存在，没有值的用空字符串""
4. messages数组至少包含2-3条关键对话
5. 如果无法判断类型，返回：{"type": "text", "content": "识别的原始文字内容"}"""
                            }
                        ]
                    }
                ],
                max_tokens=2000,
                temperature=0.3
            )
            raw_content = response.choices[0].message.content.strip()

            # 尝试提取JSON（可能被包裹在markdown代码块中）
            import re
            import json

            # 尝试从markdown代码块中提取
            json_match = re.search(r'```(?:json)?\s*\n?([\s\S]*?)\n?```', raw_content)
            if json_match:
                return json_match.group(1).strip()

            # 尝试直接解析为JSON
            try:
                parsed = json.loads(raw_content)
                return raw_content
            except:
                pass

            # 如果不是JSON格式，包装成text类型
            return json.dumps({"type": "text", "content": raw_content}, ensure_ascii=False)

        except Exception as e:
            return json.dumps({"type": "text", "content": f"识别失败: {str(e)}"}, ensure_ascii=False)

    def _process_video(self, file_content: bytes, filename: str) -> dict:
        """处理视频文件 - 暂不支持内容识别"""
        return {
            "type": "video",
            "content": "[视频文件已上传，暂不支持文字识别]",
            "filename": filename
        }


    def _process_pdf(self, file_content: bytes, filename: str) -> dict:
        """处理 PDF 文件"""
        if not HAS_PDF:
            return {
                'type': 'document',
                'content': 'PDF 处理库未安装，请安装 pypdf',
                'filename': filename
            }

        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)

            text_content = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- 第 {i+1} 页 ---\n{text}")

            full_text = "\n\n".join(text_content)

            # 如果文本太少，可能是扫描版 PDF，尝试用 OCR
            if len(full_text.strip()) < 100:
                return {
                    'type': 'document',
                    'content': "此 PDF 可能是扫描件或图片格式，暂无法提取文字。请尝试上传原始图片。",
                    'filename': filename
                }

            return {
                'type': 'document',
                'content': full_text,
                'filename': filename
            }

        except Exception as e:
            return {
                'type': 'document',
                'content': f'PDF 处理失败: {str(e)}',
                'filename': filename
            }

    def _process_word(self, file_content: bytes, filename: str) -> dict:
        """处理 Word 文档"""
        if not HAS_DOCX:
            return {
                'type': 'document',
                'content': 'Word 处理库未安装，请安装 python-docx',
                'filename': filename
            }

        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)

            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_text))

            return {
                'type': 'document',
                'content': "\n".join(text_content),
                'filename': filename
            }

        except Exception as e:
            return {
                'type': 'document',
                'content': f'Word 处理失败: {str(e)}',
                'filename': filename
            }

    def _process_excel(self, file_content: bytes, filename: str) -> dict:
        """处理 Excel 文件"""
        try:
            import openpyxl
            excel_file = io.BytesIO(file_content)
            wb = openpyxl.load_workbook(excel_file, data_only=True)

            text_content = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_content.append(f"=== 工作表: {sheet_name} ===")

                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_text):
                        text_content.append(" | ".join(row_text))

            return {
                'type': 'document',
                'content': "\n".join(text_content),
                'filename': filename
            }

        except Exception as e:
            return {
                'type': 'document',
                'content': f'Excel 处理失败: {str(e)}',
                'filename': filename
            }

    def _process_text(self, file_content: bytes, filename: str) -> dict:
        """处理纯文本文件"""
        try:
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    text = file_content.decode(encoding)
                    return {
                        'type': 'document',
                        'content': text,
                        'filename': filename
                    }
                except UnicodeDecodeError:
                    continue

            return {
                'type': 'document',
                'content': '无法识别文件编码',
                'filename': filename
            }

        except Exception as e:
            return {
                'type': 'document',
                'content': f'文本处理失败: {str(e)}',
                'filename': filename
            }


# 全局处理器实例
doc_processor = DocumentProcessor()


def process_uploaded_files(files: list) -> str:
    """
    批量处理上传的文件，返回合并后的文本

    Args:
        files: [(filename, content), ...] 文件列表

    Returns:
        str: 提取的所有文本内容
    """
    results = []

    for filename, content in files:
        result = doc_processor.process_file(content, filename)

        if result['type'] == 'unsupported':
            results.append(f"【{filename}】: {result['content']}")
        else:
            results.append(f"=== {filename} ===\n{result['content']}")

    return "\n\n".join(results)
